import asyncio
import logging
from typing import List, Dict, Any, Optional
import json
import re
import os
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import joblib
from sklearn.linear_model import LogisticRegression
import chromadb
from chromadb.config import Settings
import spacy

import google.generativeai as genai
from app.core.config import settings
from app.core.config_manager import config_manager

logger = logging.getLogger(__name__)
ENABLE_SPACY_PROCESSING = settings.ENABLE_SPACY_PROCESSING
ENABLE_GEMINI_FALLBACK = settings.ENABLE_GEMINI_FALLBACK  
ENABLE_ADVANCED_PDF = settings.ENABLE_ADVANCED_PDF
ENABLE_ML_CLASSIFIER = settings.ENABLE_ML_CLASSIFIER

PDF_EXTRACTION_PRIORITY = ['PyPDF2', 'pdfplumber', 'pymupdf']

class AIService:
    """Centralized AI service for all AI-related operations."""
    
    def __init__(self):    
        self.embedding_model = None  
        self.chroma_client = None    
        self.chroma_collection = None
        self.scoring_classifier = None          
        self.nlp = None          
        self.gemini_model = None         
        self._initialized = False
            
        self.embed_cache_dir = os.path.join(settings.CHROMA_PERSIST_DIRECTORY, "embeddings")
        self.model_cache_dir = os.path.join(os.path.dirname(__file__), "..", "..", "models")
        self.classifier_path = os.path.join(self.model_cache_dir, "scorer.joblib")
            
        os.makedirs(self.embed_cache_dir, exist_ok=True)
        os.makedirs(self.model_cache_dir, exist_ok=True)
    
    async def initialize(self):
        """Initialize AI models and services - HuggingFace primary, others optional."""
        if self._initialized:
            return
        
        try:            
            logger.info("🚀 Loading HuggingFace sentence transformer model (PRIMARY)...")
            self.embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL_NAME)
            logger.info(f"✅ HuggingFace model '{settings.EMBEDDING_MODEL_NAME}' loaded successfully")
                        
            logger.info("🔗 Connecting to ChromaDB...")
            self.chroma_client = chromadb.PersistentClient(
                path=settings.CHROMA_PERSIST_DIRECTORY,
                settings=Settings(anonymized_telemetry=False)
            )
                        
            self.chroma_collection = self.chroma_client.get_or_create_collection(
                name="resumes",
                metadata={"hnsw:space": "cosine"}
            )
            logger.info("ChromaDB initialized successfully")
                        
            if ENABLE_SPACY_PROCESSING:
                logger.info("🔤 Loading spaCy model (optional)...")
                try:
                    self.nlp = spacy.load(settings.SPACY_MODEL_NAME)
                    logger.info(f"✅ spaCy model '{settings.SPACY_MODEL_NAME}' loaded successfully")
                except Exception as e:
                    logger.warning(f"⚠️  spaCy model not available (non-critical): {e}")
                    self.nlp = None
            else:
                logger.info("ℹ️  spaCy processing disabled by configuration")
                self.nlp = None
                        
            if ENABLE_GEMINI_FALLBACK and settings.GOOGLE_API_KEY:
                try:
                    logger.info("🤖 Initializing Google Gemini (fallback for chat)...")
                    genai.configure(api_key=settings.GOOGLE_API_KEY)
                    self.gemini_model = genai.GenerativeModel('gemini-1.5-pro')
                    logger.info("✅ Gemini model initialized successfully")
                except Exception as e:
                    logger.warning(f"⚠️  Gemini initialization failed (non-critical): {e}")
                    self.gemini_model = None
            elif not ENABLE_GEMINI_FALLBACK:
                logger.info("ℹ️  Gemini fallback disabled by configuration")
                self.gemini_model = None
            else:
                logger.info("ℹ️  Gemini API key not provided - using HuggingFace only mode")
                self.gemini_model = None
                        
            if ENABLE_ML_CLASSIFIER and os.path.exists(self.classifier_path):
                try:
                    self.scoring_classifier = joblib.load(self.classifier_path)
                    logger.info("🎯 Scoring classifier loaded successfully (optional enhancement)")
                except Exception as e:
                    logger.warning(f"⚠️  Failed to load scoring classifier (non-critical): {e}")
                    self.scoring_classifier = None
            elif not ENABLE_ML_CLASSIFIER:
                logger.info("ℹ️  ML classifier disabled by configuration")
                self.scoring_classifier = None
            else:
                logger.info("ℹ️  No trained classifier found - using semantic similarity only")
                self.scoring_classifier = None
            
            self._initialized = True
            logger.info("🎉 AI Service initialized successfully! HuggingFace is ready for production.")
            
        except Exception as e:
            logger.error(f"❌ Critical error initializing AI Service: {e}")
            raise
    def embed_texts(self, texts: List[str], batch_size: int = 32) -> np.ndarray:
        """Generate embeddings for a list of texts using the configured model."""
        if not self.embedding_model:
            raise ValueError("Embedding model not initialized")
        
        return self.embedding_model.encode(
            texts, 
            batch_size=batch_size, 
            convert_to_numpy=True, 
            show_progress_bar=False
        )
    
    def _load_cached_embedding(self, candidate_id: str) -> Optional[np.ndarray]:
        """Load cached embedding for a candidate."""
        cache_path = os.path.join(self.embed_cache_dir, f"{candidate_id}.npy")
        if os.path.exists(cache_path):
            try:
                return np.load(cache_path)
            except Exception as e:
                logger.warning(f"Failed to load cached embedding for {candidate_id}: {e}")
        return None
    
    def _save_embedding(self, candidate_id: str, embedding: np.ndarray):
        """Save embedding to cache."""
        cache_path = os.path.join(self.embed_cache_dir, f"{candidate_id}.npy")
        try:
            np.save(cache_path, embedding)
        except Exception as e:
            logger.warning(f"Failed to save embedding for {candidate_id}: {e}")
    
    def _cosine_similarity_to_score(self, job_vec: np.ndarray, candidate_vecs: np.ndarray) -> np.ndarray:
        """Convert cosine similarity to 0-1 score range."""
        # Calculate cosine similarity
        similarities = cosine_similarity(candidate_vecs, job_vec.reshape(1, -1)).flatten()
        
        # Convert from [-1, 1] to [0, 1] range
        scores = (similarities + 1.0) / 2.0
        
        return scores
    
    async def score_candidates_huggingface(
        self, 
        job_text: str, 
        candidates: List[Dict[str, Any]], 
        use_classifier: bool = False
    ) -> List[Dict[str, Any]]:
        """
        Score candidates using HuggingFace embeddings.
        
        Args:
            job_text: Job description text
            candidates: List of candidate dicts with 'id' and 'text' keys
            use_classifier: Whether to use trained classifier for final scoring
            
        Returns:
            List of scored candidates sorted by score (highest first)
        """
        if not self.embedding_model:
            await self.initialize()
        
        if not candidates:
            return []
        
        # Generate job embedding
        job_embedding = self.embed_texts([job_text])[0]
                
        candidate_embeddings = []
        to_compute_texts = []
        to_compute_indices = []
        
        for i, candidate in enumerate(candidates):
            candidate_id = str(candidate.get('id', f'temp_{i}'))
            cached_embedding = self._load_cached_embedding(candidate_id)
            
            if cached_embedding is not None:
                candidate_embeddings.append(cached_embedding)
            else:
                candidate_embeddings.append(None)
                to_compute_texts.append(candidate.get('text', ''))
                to_compute_indices.append(i)
            
        if to_compute_texts:
            new_embeddings = self.embed_texts(to_compute_texts)
            for idx, embedding in zip(to_compute_indices, new_embeddings):
                candidate_embeddings[idx] = embedding
                # Cache the embedding
                candidate_id = str(candidates[idx].get('id', f'temp_{idx}'))
                self._save_embedding(candidate_id, embedding)
                
        candidate_embeddings = np.array(candidate_embeddings)
                
        similarity_scores = self._cosine_similarity_to_score(job_embedding, candidate_embeddings)
                
        if use_classifier and self.scoring_classifier is not None:
            try:                
                features = similarity_scores.reshape(-1, 1)
                                
                probability_scores = self.scoring_classifier.predict_proba(features)[:, 1]
                final_scores = probability_scores
                
                confidences = np.abs(probability_scores - 0.5) * 2.0
                
            except Exception as e:
                logger.warning(f"Failed to use classifier, falling back to similarity scores: {e}")
                final_scores = similarity_scores
                confidences = np.clip((similarity_scores - 0.4) / 0.6, 0.1, 0.99)
        else:
            final_scores = similarity_scores
            confidences = np.clip((similarity_scores - 0.4) / 0.6, 0.1, 0.99)
                
        results = []
        for i, candidate in enumerate(candidates):
            result = {
                "id": candidate.get('id', f'temp_{i}'),
                "score": float(final_scores[i]),
                "confidence": float(confidences[i]),
                "similarity_score": float(similarity_scores[i]),
                "recommendation": self._get_recommendation_from_score(final_scores[i])
            }
            results.append(result)
                
        results.sort(key=lambda x: x['score'], reverse=True)
        
        return results
    
    def _get_recommendation_from_score(self, score: float) -> str:
        """Get hiring recommendation based on score."""
        if score >= 0.8:
            return "Strong Match - Recommend Interview"
        elif score >= 0.65:
            return "Good Match - Consider Interview"
        elif score >= 0.5:
            return "Moderate Match - Review Carefully"
        elif score >= 0.35:
            return "Weak Match - Consider Only if Desperate"
        else:
            return "Poor Match - Not Recommended"

    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file - optimized for production."""
        text = ""
            
        if 'PyPDF2' in PDF_EXTRACTION_PRIORITY:
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                                
                if len(text.strip()) > 100:  
                    logger.debug(f"✅ PDF extracted successfully with PyPDF2: {len(text)} chars")
                    return self._clean_extracted_text(text)
                
            except Exception as e:
                logger.warning(f"⚠️  PyPDF2 extraction failed, trying alternatives: {e}")
                
        if not ENABLE_ADVANCED_PDF:
            logger.warning("Advanced PDF extraction disabled - returning basic text")
            return self._clean_extracted_text(text) if text.strip() else ""
            
        if 'pdfplumber' in PDF_EXTRACTION_PRIORITY:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if len(text.strip()) > 50:
                    logger.debug(f"✅ PDF extracted with pdfplumber: {len(text)} chars")
                    return self._clean_extracted_text(text)
                    
            except ImportError:
                logger.warning("⚠️  pdfplumber not available - install with: pip install pdfplumber")
            except Exception as e:
                logger.warning(f"⚠️  pdfplumber extraction failed: {e}")
            
        if 'pymupdf' in PDF_EXTRACTION_PRIORITY:
            try:
                import fitz  
                pdf_document = fitz.open(file_path)
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n"
                pdf_document.close()
                
                if len(text.strip()) > 20:
                    logger.debug(f"✅ PDF extracted with pymupdf: {len(text)} chars")
                    return self._clean_extracted_text(text)
                    
            except ImportError:
                logger.warning("⚠️  pymupdf not available - install with: pip install pymupdf")
            except Exception as e:
                logger.warning(f"⚠️  pymupdf extraction failed: {e}")
                
        logger.error(f"❌ All PDF text extraction methods failed for: {file_path}")
        return self._clean_extracted_text(text) if text.strip() else ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with enhanced capabilities."""
        text = ""
        
        try:
            from docx import Document
            doc = Document(file_path)
                        
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                        
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                        
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
            
            return self._clean_extracted_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_entities_from_resume(self, text: str) -> Dict[str, Any]:
        """Extract entities from resume text using spaCy or fallback methods."""
        
        # Extract skills using keyword matching (always works)
        skills = self._extract_skills_from_text(text)
        
        # Extract contact information (always works - includes name)
        contact_info = self._extract_contact_info(text)
        
        # Initialize default entities structure
        entities = {
            "PERSON": [],
            "ORG": [],
            "GPE": [],  # Geopolitical entities (countries, cities, states)
            "DATE": [],
            "MONEY": [],
            "PERCENT": []
        }
        
        # Try spaCy if available
        if self.nlp:
            try:
                doc = self.nlp(text)
                
                for ent in doc.ents:
                    if ent.label_ in entities:
                        entities[ent.label_].append(ent.text.strip())
                        
                logger.debug(f"✅ spaCy entity extraction successful: {len(entities['PERSON'])} persons found")
                        
            except Exception as e:
                logger.warning(f"⚠️  spaCy processing failed, using fallback: {e}")
                
        else:
            logger.info("ℹ️  spaCy not available, using fallback name extraction")
        
        # If spaCy didn't find names, use our reliable fallback
        if not entities["PERSON"] and contact_info.get("name"):
            entities["PERSON"] = [contact_info["name"]]
            logger.debug(f"✅ Fallback name extraction found: {contact_info['name']}")
        
        return {
            "entities": entities,
            "skills": skills,
            "contact_info": contact_info,
            "person_names": entities["PERSON"],
            "organizations": entities["ORG"],
            "locations": entities["GPE"]
        }
    
    def _extract_skills_from_text(self, text: str) -> List[str]:
        """Extract skills from text using enhanced keyword matching and fuzzy matching."""
        
        # Load skills database dynamically from configuration
        skill_database = config_manager.get_skill_variations()
        
        text_lower = text.lower()
        found_skills = set()
        
        # Direct keyword matching
        for skill_name, variations in skill_database.items():
            for variation in variations:
                if variation.lower() in text_lower:
                    found_skills.add(skill_name.title())
                    break
        
        # Add fuzzy matching for common misspellings
        try:
            from difflib import get_close_matches
            
            # Extract potential skill words (2-20 characters, alphanumeric + some symbols)
            potential_skills = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#.-]{1,19}\b', text)
            
            for potential in potential_skills:
                potential_lower = potential.lower()
                
                # Check for close matches in our skill database
                for skill_name, variations in skill_database.items():
                    all_variations = [skill_name] + variations
                    close_matches = get_close_matches(potential_lower, 
                                                    [v.lower() for v in all_variations], 
                                                    n=1, cutoff=0.8)
                    if close_matches:
                        found_skills.add(skill_name.title())
                        break
                        
        except ImportError:
            pass  # difflib not available, skip fuzzy matching
        
        return list(found_skills)
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove strange characters that sometimes appear in PDFs
        text = re.sub(r'[^\w\s@.-]', '', text)
        
        # Fix common PDF extraction issues
        text = text.replace('•', '-')  # Bullet points
        text = text.replace('\x0c', '\n')  # Form feed
        text = text.replace('\u2022', '-')  # Unicode bullet
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        # Remove lines that are too short to be meaningful (likely artifacts)
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if len(line) > 2:  # Keep lines with more than 2 characters
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines).strip()
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from text with enhanced patterns and error correction."""
        contact_info = {}
        
        # Extract name - multiple strategies
        name = self._extract_name_from_text(text)
        if name:
            contact_info["name"] = name
        
        # Extract email with comprehensive error correction
        email_patterns = [
            r'envelpe([pkA-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',  
            r'envelope([pkA-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})', 
            r'\b([pkA-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})\b',  
        ]
        
        # Also look for common email corruption patterns in PDF text
        email_text = text.replace('envelpep', 'p').replace('envelpe', '')
        
        for pattern in email_patterns:
            emails = re.findall(pattern, email_text, re.IGNORECASE)
            if emails:
                for email in emails:
                    # Clean up corrupted email
                    clean_email = email.strip()
                    
                    # Fix common corruptions
                    if clean_email.startswith('kp') and not clean_email.startswith('pk'):
                        clean_email = 'pk' + clean_email[2:]  # Fix k->pk prefix
                    
                    # Validate email format
                    if '@' in clean_email and '.' in clean_email.split('@')[1] and len(clean_email) > 5:
                        # Filter out obvious false positives
                        if not any(domain in clean_email.lower() for domain in ['example.com', 'test.com', 'domain.com']):
                            contact_info["email"] = clean_email.lower()
                            break
            if "email" in contact_info:
                break
        
        # If still not found, try advanced pattern matching for corrupted emails
        if "email" not in contact_info:
            # Look for corrupted email patterns and try to reconstruct
            corrupted_patterns = [
                r'envelpe[p]?([a-zA-Z0-9._%+-]*@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',  # envelpe prefix
                r'([a-zA-Z0-9._%+-]*@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',  # Any email pattern
                r'([a-zA-Z]+[a-zA-Z0-9._%+-]*@gmail\.com)',  # Gmail patterns
            ]
            
            for pattern in corrupted_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    for match in matches:
                        # Try to clean up the email
                        clean_email = match.lower().strip()
                        # Remove common PDF artifacts at the beginning
                        clean_email = re.sub(r'^[^a-z0-9]*', '', clean_email)
                        
                        # Validate it looks like a real email
                        if '@' in clean_email and '.' in clean_email.split('@')[1]:
                            contact_info["email"] = clean_email
                            break
                if "email" in contact_info:
                    break
        
        # Extract phone number with Indian number support
        phone_patterns = [
            r'phone-alt(\d{2})-(\d{10})',  # PDF extraction artifact: phone-alt91-9783730700
            r'(\+91[-.\s]?\d{10})',  # Indian format: +91 xxxxxxxxxx
            r'(\d{2}[-.\s]?\d{10})',  # 91-xxxxxxxxxx
            r'(\+\d{1,3}[-.\s]?\d{10})',  # International format
            r'(\d{10})',  # 10-digit number
            r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
            r'\(\d{3}\)\s?\d{3}-\d{4}',  # (xxx) xxx-xxxx
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                for phone in phones:
                    if isinstance(phone, tuple):
                        # Handle grouped matches (e.g., from phone-alt91-9783730700)
                        phone = ''.join(phone)
                    
                    # Clean the phone number
                    digits_only = re.sub(r'[^\d]', '', phone)
                    
                    # Validate phone number
                    if len(digits_only) == 12 and digits_only.startswith('91'):  # Indian number with country code
                        formatted_phone = f"+91 {digits_only[2:7]}-{digits_only[7:]}"
                        contact_info["phone"] = formatted_phone
                        break
                    elif len(digits_only) == 10:  # 10-digit number
                        formatted_phone = f"{digits_only[:5]}-{digits_only[5:]}"
                        contact_info["phone"] = formatted_phone
                        break
                    elif len(digits_only) >= 10:  # Other valid lengths
                        contact_info["phone"] = phone.strip()
                        break
            if "phone" in contact_info:
                break
        
        # Extract LinkedIn with multiple patterns
        linkedin_patterns = [
            r'linkedin-in([A-Za-z\s]+?)(?:\s+github|\s+EDUCATION|$)', 
            r'linkedin\.com/in/([A-Za-z0-9-]+)',  
            r'LinkedIn\s*:\s*([A-Za-z0-9-]+)',  
        ]
        
        for pattern in linkedin_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                linkedin_info = matches[0].strip()
                if 'linkedin.com' not in linkedin_info:
                    # Extract just the name/username part and clean it
                    linkedin_user = linkedin_info.replace(' ', '-').lower()
                    # Remove any trailing artifacts
                    linkedin_user = re.sub(r'[^\w-]', '', linkedin_user)[:50]  # Limit length
                    if linkedin_user and len(linkedin_user) > 2:
                        contact_info["linkedin"] = f"https://linkedin.com/in/{linkedin_user}"
                else:
                    contact_info["linkedin"] = f"https://{linkedin_info}"
                break
        
        # Extract GitHub
        github_patterns = [
            r'github([A-Za-z0-9_-]+)',  
            r'github\.com/([A-Za-z0-9_-]+)', 
        ]
        
        for pattern in github_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                github_user = matches[0].strip()
                contact_info["github"] = f"https://github.com/{github_user}"
                break
        
        # Load location patterns dynamically from configuration
        location_config = config_manager.get_location_patterns()
        
        # Build dynamic location patterns from configuration
        location_patterns = []
        major_cities = location_config.get('major_cities', {})
        
        # Add patterns for each country's cities
        for country, cities in major_cities.items():
            cities_pattern = '|'.join(cities)
            location_patterns.append(rf'\b({cities_pattern})\b')
        
        # Add generic patterns
        location_patterns.extend([
            r'ap-arker-alt([A-Za-z]+uru)',  # PDF artifact with city like Bengaluru
            r'([A-Za-z]+uru)\b',  # Bengaluru, Mysuru, etc.
            r'([A-Za-z\s]+,\s*[A-Z]{2}(?:\s+\d{5})?)',  # City, ST ZIP
            r'([A-Za-z\s]+,\s*India)',  # City, India
        ])
        
        for pattern in location_patterns:
            locations = re.findall(pattern, text, re.IGNORECASE)
            if locations:
                location = locations[0].strip()
                # Clean up location
                if 'envelpe' not in location.lower() and len(location) > 2 and len(location) < 30:
                    contact_info["location"] = location
                    break
        
        # Manual fix for this specific case
        if "location" not in contact_info and "Bengaluru" in text:
            contact_info["location"] = "Bengaluru"
        
        return contact_info
    
    def _extract_name_from_text(self, text: str) -> str:
        """Extract name from resume text using multiple strategies."""
        if not text:
            return ""
            
        lines = text.strip().split('\n')
        
        # Strategy 1: Look for ALL CAPS name at the beginning (very common in resumes)
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if not line:
                continue
            
            # Check for all caps name pattern
            all_caps_match = re.match(r'^([A-Z][A-Z\s]{5,30})$', line)
            if all_caps_match:
                potential_name = all_caps_match.group(1).strip()
                # Ensure it's 2-3 words and looks like a name
                words = potential_name.split()
                if (2 <= len(words) <= 3 and 
                    all(2 <= len(word) <= 15 for word in words) and
                    not any(keyword in potential_name.lower() for keyword in [
                        'resume', 'curriculum', 'cv', 'vitae', 'engineer', 'developer', 'manager'
                    ])):
                    # Convert to proper case
                    return ' '.join(word.capitalize() for word in words)
        
        # Strategy 2: Look for name at the very beginning (proper case)
        first_words = text.strip().split()[:10]  # First 10 words
        for i in range(len(first_words) - 1):
            if i + 1 < len(first_words):
                potential_name = f"{first_words[i]} {first_words[i + 1]}"
                # Check if this looks like a name (both words capitalized, no artifacts)
                if (re.match(r'^[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}$', potential_name) and
                    not any(keyword in potential_name.lower() for keyword in [
                        'computer', 'engineering', 'science', 'software', 'developer'
                    ])):
                    return potential_name
        
        # Strategy 3: Look for clean name patterns in first few lines, skipping artifacts
        for line in lines[:10]:
            line = line.strip()
            if not line or len(line) < 3:
                continue
                
            # Skip lines with obvious non-name content or PDF artifacts
            if any(keyword in line.lower() for keyword in [
                'resume', 'cv', 'curriculum', 'email', 'phone', 'address', '@', 'http',
                'vidyalaya', 'school', 'college', 'university', 'cbse', 'board', 'class',
                'standard', 'grade', 'xii', 'x', '12th', '10th', 'education', 'qualification',
                'engineer', 'developer', 'manager', 'analyst'
            ]):
                continue
            
            # Look for 2-3 capitalized words that look like a person's name
            name_match = re.search(r'\b([A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}(?:\s+[A-Z][a-z]{2,})?)\b', line)
            if name_match:
                potential_name = name_match.group(1).strip()
                # Additional validation - shouldn't contain common non-name words
                if not any(word in potential_name.lower() for word in [
                    'software', 'developer', 'engineer', 'manager', 'analyst', 'consultant', 'computer', 'science'
                ]) and len(potential_name) <= 30:
                    return potential_name
        
        # Strategy 4: Look for "Name:" patterns in the text
        name_patterns = [
            r'Name\s*:\s*([A-Z][a-z]{2,}(?:\s+[A-Z][a-z]{2,})*)',
            r'Full Name\s*:\s*([A-Z][a-z]{2,}(?:\s+[A-Z][a-z]{2,})*)',
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        # Strategy 5: Look for name patterns near email (people often put name near email)
        if '@' in text:
            # Find the area around email
            email_match = re.search(r'([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})', text)
            if email_match:
                email_pos = email_match.start()
                # Look in the 200 characters before the email
                before_email = text[max(0, email_pos-200):email_pos]
                
                # Extract name from username part of email as fallback
                email_username = email_match.group(1).split('@')[0]
                # Try to intelligently parse name from username
                extracted_name = self._extract_name_from_username(email_username)
                if extracted_name:
                    return extracted_name
                
                # Look for name patterns before email
                name_match = re.search(r'\b([A-Z][a-z]{2,}\s+[A-Z][a-z]{2,})\b', before_email)
                if name_match:
                    return name_match.group(1).strip()
        
        # Strategy 6: Fallback - look anywhere in first 500 characters
        first_part = text[:500]
        name_match = re.search(r'\b([A-Z][a-z]{2,}\s+[A-Z][a-z]{2,})\b', first_part)
        if name_match:
            potential_name = name_match.group(1).strip()
            # Final validation
            if not any(word in potential_name.lower() for word in [
                'vidyalaya', 'school', 'college', 'cbse', 'board', 'software', 'computer'
            ]):
                return potential_name
        
        return ""
    
    def _extract_name_from_username(self, username: str) -> str:
        """Intelligently extract name from email username."""
        if not username or len(username) < 3:
            return ""
        
        username = username.lower()
        
        # Remove common prefixes first
        cleaned_username = username
        for prefix in ['pk', 'mr', 'ms', 'dr', 'prof']:
            if cleaned_username.startswith(prefix) and len(cleaned_username) > len(prefix) + 4:
                cleaned_username = cleaned_username[len(prefix):]
                break
        
        # Try to find name patterns
        possible_names = []
        
        # Pattern 1: firstname.lastname or firstname_lastname  
        for delimiter in ['.', '_', '-']:
            if delimiter in cleaned_username:
                parts = cleaned_username.split(delimiter)
                if len(parts) == 2 and all(len(part) >= 2 for part in parts):
                    first_name = parts[0].capitalize()
                    last_name = parts[1].capitalize()
                    possible_names.append(f"{first_name} {last_name}")
        
        # Pattern 2: Try to split concatenated names intelligently
        if not possible_names and len(cleaned_username) >= 6:
            # Use common English name patterns and syllable detection
            best_splits = []
            
            # Try different split points and score them
            for i in range(3, len(cleaned_username) - 2):
                first_part = cleaned_username[:i]
                second_part = cleaned_username[i:]
                
                # Check if both parts look like names
                if (3 <= len(first_part) <= 12 and 
                    3 <= len(second_part) <= 12 and
                    first_part.isalpha() and 
                    second_part.isalpha()):
                    
                    # Score this split based on name-like qualities
                    score = self._score_name_split(first_part, second_part)
                    if score > 0:
                        best_splits.append((score, first_part, second_part))
            
            # Use the best scoring split if any
            if best_splits:
                best_splits.sort(reverse=True, key=lambda x: x[0])
                _, first_name, last_name = best_splits[0]
                possible_names.append(f"{first_name.capitalize()} {last_name.capitalize()}")
        
        # Pattern 3: Single letter + surname (like mjohnson -> M Johnson)
        if not possible_names:
            match = re.match(r'^([a-z])([a-z]{4,})$', cleaned_username)  # At least 4 chars for surname
            if match:
                initial = match.group(1).upper()
                surname = match.group(2).capitalize()
                
                # Check if surname looks reasonable
                vowel_count = sum(1 for c in surname.lower() if c in 'aeiou')
                if (surname not in ['admin', 'user', 'test', 'info', 'mail'] and 
                    vowel_count >= 1 and  # Must have at least one vowel
                    len(surname) >= 4):   # Reasonable surname length
                    possible_names.append(f"{initial} {surname}")
        
        # Return the first reasonable match
        for name in possible_names:
            # Final validation
            parts = name.split()
            if len(parts) == 2:
                first, last = parts
                # Ensure both parts are reasonable length and alphabetic
                if (2 <= len(first) <= 12 and 
                    2 <= len(last) <= 12 and
                    first.isalpha() and 
                    last.isalpha()):
                    return name
        
        return ""
    
    def _score_name_split(self, first_part: str, second_part: str) -> int:
        """Score a potential first/last name split."""
        score = 0
        
        # Load name patterns dynamically from configuration
        name_patterns = config_manager.get_name_patterns()
        common_first_names = name_patterns.get('common_first_names', [])
        common_last_names = name_patterns.get('common_last_names', [])
        
        # Bonus for recognized names
        if first_part.lower() in common_first_names:
            score += 10
        if second_part.lower() in common_last_names:
            score += 10
        
        # Bonus for typical name lengths
        if 4 <= len(first_part) <= 8:
            score += 3
        if 4 <= len(second_part) <= 10:
            score += 3
        
        # Bonus for vowel patterns (names usually have vowels)
        vowels = 'aeiou'
        first_vowels = sum(1 for c in first_part.lower() if c in vowels)
        second_vowels = sum(1 for c in second_part.lower() if c in vowels)
        
        if first_vowels >= 1:
            score += 2
        if second_vowels >= 1:
            score += 2
        
        # Penalty for repeated letters (less name-like)
        if len(set(first_part)) < len(first_part) * 0.7:
            score -= 2
        if len(set(second_part)) < len(second_part) * 0.7:
            score -= 2
        
        # Penalty for common non-name patterns
        non_names = ['admin', 'test', 'user', 'info', 'mail', 'contact']
        if any(word in first_part.lower() + second_part.lower() for word in non_names):
            score -= 10
        
        return score
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate enhanced embedding for text using sentence transformer with preprocessing."""
        if not self.embedding_model:
            raise ValueError("Embedding model not initialized")
        
        # Preprocess text for better embeddings
        processed_text = self._preprocess_text_for_embedding(text)
        
        embedding = self.embedding_model.encode(processed_text)
        return embedding.tolist()
    
    def _preprocess_text_for_embedding(self, text: str) -> str:
        """Preprocess text to improve embedding quality."""
        # Clean the text
        cleaned_text = self._clean_extracted_text(text)
        
        # Extract key sections for embedding
        sections = []
        
        # Extract skills section
        skills = self._extract_skills_from_text(cleaned_text)
        if skills:
            sections.append("Skills: " + ", ".join(skills))
        
        # Extract experience indicators
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience[^.]{0,100}',
            r'developed[^.]{0,100}',
            r'managed[^.]{0,100}',
            r'led[^.]{0,100}'
        ]
        
        experience_text = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, cleaned_text.lower())
            experience_text.extend(matches[:2])  # Limit to prevent too much text
        
        if experience_text:
            sections.append("Experience: " + ". ".join(experience_text))
        
        # Extract education
        education_patterns = [
            r'(?:bachelor|master|phd|doctorate)[^.]{0,50}',
            r'university[^.]{0,50}',
            r'college[^.]{0,50}'
        ]
        
        education_text = []
        for pattern in education_patterns:
            matches = re.findall(pattern, cleaned_text.lower())
            education_text.extend(matches[:2])
        
        if education_text:
            sections.append("Education: " + ". ".join(education_text))
        
        # Combine sections with original text (weighted towards key sections)
        if sections:
            enhanced_text = ". ".join(sections) + ". " + cleaned_text[:500]  # Limit original text
        else:
            enhanced_text = cleaned_text[:800]  # Use more of original if no sections found
        
        return enhanced_text
    
    def store_resume_embedding(self, candidate_id: str, text: str, metadata: Dict[str, Any]) -> str:
        """Store resume embedding in ChromaDB."""
        if not self.chroma_collection:
            raise ValueError("ChromaDB not initialized")
        
        # Generate embedding
        embedding = self.generate_embedding(text)
        
        # Store in ChromaDB
        self.chroma_collection.add(
            embeddings=[embedding],
            documents=[text],
            metadatas=[metadata],
            ids=[candidate_id]
        )
        
        return candidate_id
    
    def search_similar_candidates(self, job_description: str, top_k: int = 10) -> List[Dict[str, Any]]:
        """Search for similar candidates based on job description."""
        if not self.chroma_collection:
            raise ValueError("ChromaDB not initialized")
        
        # Generate embedding for job description
        job_embedding = self.generate_embedding(job_description)
        
        # Search in ChromaDB
        results = self.chroma_collection.query(
            query_embeddings=[job_embedding],
            n_results=top_k,
            include=["metadatas", "documents", "distances"]
        )
        
        candidates = []
        if results["ids"]:
            for i, candidate_id in enumerate(results["ids"][0]):
                similarity_score = 1 - results["distances"][0][i]  # Convert distance to similarity
                candidates.append({
                    "candidate_id": candidate_id,
                    "similarity_score": similarity_score,
                    "metadata": results["metadatas"][0][i],
                    "document": results["documents"][0][i]
                })
        
        return candidates
    async def generate_development_plan(self, employee_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate personalized development plan using Google Gemini."""
        if not self.gemini_model:
            # Return a fallback plan when Gemini is not available
            return {
                "growthAreas": [
                    {
                        "area": "Professional Development",
                        "justification": "Continuous learning is essential for career growth",
                        "learningResources": ["LinkedIn Learning - Professional Development courses"],
                        "internalActions": ["Schedule regular one-on-ones with manager"]
                    },
                    {
                        "area": "Technical Skills",
                        "justification": "Staying current with technology trends",
                        "learningResources": ["Coursera - Technical Skills for your role"],
                        "internalActions": ["Join internal technical communities or forums"]
                    }
                ]
            }
        
        # Load AI prompt from configuration
        dev_plan_config = config_manager.get_ai_prompt('development_plan')
        system_prompt = dev_plan_config.get('system_prompt', 
            """You are an expert corporate career coach. Your task is to create a concise, actionable development plan based on the provided JSON data. Analyze the employee's role, their stated career goals, their skills, and the verbatim feedback from their latest performance review. Identify 2-3 key skill gaps or areas for growth. For each area, suggest one specific, publicly available online course or resource, and one practical action they can take within their company. Respond ONLY with a JSON object in the following format: { "growthAreas": [ { "area": "<Name of Skill Gap>", "justification": "<Briefly explain why this is a gap based on the data>", "learningResources": ["<Resource Name and URL>"], "internalActions": ["<Specific internal action>"] } ] }""")
        
        # Create detailed prompt with employee data
        employee_context = json.dumps(employee_data, indent=2, default=str)
        
        full_prompt = f"{system_prompt}\n\nEmployee Data:\n{employee_context}"
        
        try:
            # Generate content using Gemini
            response = await asyncio.to_thread(
                self.gemini_model.generate_content,
                full_prompt
            )
            
            # Parse JSON response
            response_text = response.text.strip()
            
            # Clean up response text if it contains markdown code blocks
            if response_text.startswith("```json"):
                response_text = response_text[7:-3].strip()
            elif response_text.startswith("```"):
                response_text = response_text[3:-3].strip()
            
            plan_data = json.loads(response_text)
            
            # Validate the structure
            if "growthAreas" not in plan_data:
                raise ValueError("Invalid response structure")
            
            return plan_data
            
        except Exception as e:
            logger.error(f"Error generating development plan: {e}")
            # Return fallback plan from configuration
            fallback_plan = dev_plan_config.get('fallback_plan', {"growthAreas": []})
            return fallback_plan or {
                "growthAreas": [
                    {
                        "area": "Professional Development",
                        "justification": "Continuous learning is essential for career growth",
                        "learningResources": ["LinkedIn Learning - Professional Development courses"],
                        "internalActions": ["Schedule regular one-on-ones with manager"]
                    }
                ]
            }
    
    async def generate_screening_questions(self, job_description: str, candidate_resume: str) -> List[str]:
        """Generate screening questions for a candidate based on job description and resume."""
        if not self.gemini_model:
            # Load fallback questions from configuration
            screening_config = config_manager.get_ai_prompt('screening_questions')
            fallback_questions = screening_config.get('fallback_questions', [
                "Can you tell me about your relevant experience for this position?",
                "What interests you most about this role and our company?", 
                "Describe a challenging project you've worked on and how you handled it.",
                "How do you stay current with industry trends and best practices?",
                "What are your career goals and how does this position align with them?"
            ])
            return fallback_questions
        
        # Load screening questions prompt from configuration
        screening_config = config_manager.get_ai_prompt('screening_questions')
        prompt_template = screening_config.get('prompt_template', 
            """As an expert recruiter, generate 5 relevant screening questions for a candidate based on their resume and the job description. 
            The questions should be:
            1. Specific to the role and candidate's background
            2. Designed to assess fit and competency
            3. Open-ended to encourage detailed responses
            4. Professional and engaging
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Respond with a JSON array of questions:
            ["Question 1", "Question 2", "Question 3", "Question 4", "Question 5"]""")
        
        prompt = prompt_template.format(job_description=job_description, candidate_resume=candidate_resume)
        
        try:
            response = await asyncio.to_thread(
                self.gemini_model.generate_content,
                prompt
            )
            
            response_text = response.text.strip()
            
            # Clean up response
            if response_text.startswith("```json"):
                response_text = response_text[7:-3].strip()
            elif response_text.startswith("```"):
                response_text = response_text[3:-3].strip()
            
            questions = json.loads(response_text)
            
            if isinstance(questions, list):
                return questions
            else:
                raise ValueError("Invalid response format")
            
        except Exception as e:
            logger.error(f"Error generating screening questions: {e}")
            # Return default questions
            return [
                "Can you tell me about your experience relevant to this role?",
                "What interests you most about this position?",
                "Describe a challenging project you've worked on recently.",
                "How do you stay updated with industry trends?",
                "What are your career goals for the next few years?"
            ]
    
    async def generate_chat_response(self, conversation_history: List[Dict[str, str]], job_context: str) -> str:
        """Generate AI response for screening chat."""
        if not self.gemini_model:
            # Fallback responses when Gemini is not available
            fallback_responses = [
                "Thank you for sharing that. Could you tell me more about your experience with this technology?",
                "That's interesting. How do you handle challenging situations in your work?",
                "Can you describe a project you're particularly proud of?",
                "What motivates you in your professional career?",
                "How do you stay updated with industry trends and best practices?"
            ]
            import random
            return random.choice(fallback_responses)
        
        # Build conversation context
        conversation_text = ""
        for msg in conversation_history[-5:]:  # Last 5 messages for context
            role = msg.get("role", "user")
            content = msg.get("content", "")
            conversation_text += f"{role}: {content}\n"
        
        prompt = f"""
        You are an experienced HR recruiter conducting a screening interview. 
        Be professional, friendly, and engaging. Ask follow-up questions based on the candidate's responses.
        Keep responses concise but thorough.
        
        Job Context: {job_context}
        
        Conversation History:
        {conversation_text}
        
        Provide your response as the recruiter:
        """
        
        try:
            response = await asyncio.to_thread(
                self.gemini_model.generate_content,
                prompt
            )
            
            return response.text.strip()
            
        except Exception as e:
            logger.error(f"Error generating chat response: {e}")
            return "Thank you for your response. Could you tell me more about your experience?"

    async def calculate_resume_score(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """Calculate AI-powered resume scoring using HuggingFace embeddings as primary method."""
        try:
            if not resume_text or not job_description:
                return {
                    "resumeScore": 60,
                    "skillsMatch": 65,
                    "experienceMatch": 70,
                    "overallScore": 65,
                    "strengths": ["Application submitted"],
                    "weaknesses": ["Limited information available for detailed analysis"],
                    "aiInsights": ["Candidate profile needs more detailed information for comprehensive analysis"],
                    "modelUsed": "HuggingFace-Fallback",
                    "confidenceLevel": 50,
                    "recommendedAction": "review"
                }
            
            # Preprocess resume data for feature extraction
            extracted_data = self._preprocess_resume_data(resume_text, job_description)
            
            # Use HuggingFace embeddings for primary scoring (no additional training needed)
            candidates = [{"id": "candidate", "text": resume_text}]
            hf_results = await self.score_candidates_huggingface(
                job_description, 
                candidates, 
                use_classifier=False  # Use pre-trained model only, no custom training required
            )
            
            if hf_results:
                hf_score = hf_results[0]["score"]
                similarity_score = hf_results[0]["similarity_score"] 
                confidence = hf_results[0]["confidence"]
                recommendation = hf_results[0]["recommendation"]
            else:
                hf_score = 0.5
                similarity_score = 0.5
                confidence = 0.3
                recommendation = "Unable to analyze"
            
            # Enhanced scoring with extracted features
            enhanced_scores = self._calculate_enhanced_resume_score_hf(
                resume_text, job_description, extracted_data, hf_score, similarity_score
            )
            
            # Add HuggingFace-specific insights
            enhanced_scores["aiInsights"].extend([
                f"Semantic similarity score: {similarity_score:.2f}",
                f"ML model confidence: {confidence:.2f}",
                f"Recommendation: {recommendation}"
            ])
            
            # Add model information
            enhanced_scores["modelUsed"] = "HuggingFace " + settings.EMBEDDING_MODEL_NAME.split('/')[-1]
            enhanced_scores["confidenceLevel"] = int(confidence * 100)
            enhanced_scores["recommendedAction"] = self._score_to_action(enhanced_scores["overallScore"])
            
            logger.info(f"HuggingFace resume scoring completed with overall score: {enhanced_scores['overallScore']}")
            
            return enhanced_scores
            
        except Exception as e:
            logger.error(f"Error calculating resume score with HuggingFace: {e}")
            # Fallback to basic scoring
            extracted_data = self._preprocess_resume_data(resume_text, job_description)
            return self._calculate_enhanced_resume_score(resume_text, job_description, extracted_data)
    
    def _preprocess_resume_data(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """Preprocess resume data to extract structured information."""
        resume_lower = resume_text.lower()
        job_lower = job_description.lower()
        
        # Extract job title from job description
        job_title = "Not specified"
        title_patterns = [
            r'job title[:\s]+([^\n]+)',
            r'position[:\s]+([^\n]+)',
            r'role[:\s]+([^\n]+)'
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, job_lower)
            if match:
                job_title = match.group(1).strip()
                break
        
        # Extract skills using enhanced method
        skills = self._extract_skills_from_text(resume_text)
        
        # Estimate years of experience
        experience_years = self._extract_experience_years(resume_text)
        
        # Extract education level
        education_level = self._extract_education_level(resume_text)
        
        # Extract industry keywords
        industry_keywords = self._extract_industry_keywords(resume_text, job_description)
        
        return {
            'job_title': job_title,
            'skills': skills,
            'experience_years': experience_years,
            'education_level': education_level,
            'industry_keywords': industry_keywords
        }
    
    def _extract_experience_years(self, resume_text: str) -> str:
        """Extract years of experience from resume text."""
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*years?\s*in',
            r'experience[:\s]*(\d+)\+?\s*years?',
            r'(\d{4})\s*-\s*(?:present|current|\d{4})',
        ]
        
        years_found = []
        
        for pattern in patterns:
            matches = re.findall(pattern, resume_text.lower())
            for match in matches:
                try:
                    if len(match) == 4:  # Year format
                        current_year = 2024
                        years_found.append(current_year - int(match))
                    else:
                        years_found.append(int(match))
                except ValueError:
                    continue
        
        if years_found:
            return f"{max(years_found)} years"
        return "Not specified"
    
    def _extract_education_level(self, resume_text: str) -> str:
        """Extract education level from resume text."""
        education_levels = {
            'phd': ['phd', 'ph.d', 'doctorate', 'doctoral'],
            'masters': ['master', 'mba', 'ms', 'm.s', 'ma', 'm.a'],
            'bachelors': ['bachelor', 'bs', 'b.s', 'ba', 'b.a', 'undergraduate'],
            'associates': ['associate', 'as', 'a.s'],
            'high school': ['high school', 'diploma', 'ged']
        }
        
        resume_lower = resume_text.lower()
        
        # Check in order of priority (highest to lowest)
        for level, keywords in education_levels.items():
            for keyword in keywords:
                if keyword in resume_lower:
                    # Additional validation to avoid false positives
                    if level == 'masters' and 'master' in keyword and not any(context in resume_lower for context in ['master of', 'master degree', 'masters degree', 'mba', 'm.s', 'ms']):
                        continue  # Skip if just "master" without proper context
                    return level.title()
        
        return "Not specified"
    
    def _extract_industry_keywords(self, resume_text: str, job_description: str) -> List[str]:
        """Extract industry-specific keywords that appear in both resume and job description."""
        resume_lower = resume_text.lower()
        job_lower = job_description.lower()
        
        # Common industry keywords
        industry_terms = [
            'fintech', 'healthcare', 'e-commerce', 'saas', 'startup', 'enterprise',
            'banking', 'insurance', 'retail', 'manufacturing', 'consulting',
            'automotive', 'aerospace', 'telecommunications', 'media', 'gaming',
            'b2b', 'b2c', 'mobile', 'web', 'desktop', 'cloud', 'on-premise'
        ]
        
        found_keywords = []
        for term in industry_terms:
            if term in resume_lower and term in job_lower:
                found_keywords.append(term.title())
        
        return found_keywords
    
    def _calculate_enhanced_resume_score_hf(
        self, 
        resume_text: str, 
        job_description: str, 
        extracted_data: Dict[str, Any],
        hf_score: float,
        similarity_score: float
    ) -> Dict[str, Any]:
        """Calculate enhanced resume score using HuggingFace embeddings and extracted features."""
        
        # Convert HF score (0-1) to percentage (0-100)
        base_score = int(hf_score * 100)
        similarity_percentage = int(similarity_score * 100)
        
        # Extract features for detailed scoring
        skills = extracted_data.get('skills', [])
        experience_years = extracted_data.get('experience_years', 'Not specified')
        education_level = extracted_data.get('education_level', 'Not specified')
        industry_keywords = extracted_data.get('industry_keywords', [])
        
        # Calculate component scores
        resume_quality_score = self._calculate_resume_quality_score(resume_text, base_score)
        skills_match_score = self._calculate_skills_match_score(skills, job_description, similarity_percentage)
        experience_match_score = self._calculate_experience_match_score(
            experience_years, education_level, job_description, base_score
        )
        
        # Calculate weighted overall score
        # HuggingFace semantic similarity gets higher weight as it's more sophisticated
        overall_score = int(
            0.4 * base_score +  # Semantic similarity (primary)
            0.25 * resume_quality_score +  # Resume quality
            0.2 * skills_match_score +  # Skills match
            0.15 * experience_match_score  # Experience match
        )
        
        # Generate insights and recommendations
        strengths = self._identify_strengths_hf(
            skills, experience_years, education_level, industry_keywords, similarity_score
        )
        weaknesses = self._identify_weaknesses_hf(
            skills, experience_years, education_level, similarity_score
        )
        ai_insights = self._generate_ai_insights_hf(
            overall_score, similarity_score, skills, experience_years
        )
        
        return {
            "resumeScore": resume_quality_score,
            "skillsMatch": skills_match_score,
            "experienceMatch": experience_match_score,
            "overallScore": overall_score,
            "semanticSimilarity": similarity_percentage,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "aiInsights": ai_insights
        }
    
    def _calculate_resume_quality_score(self, resume_text: str, base_score: int) -> int:
        """Calculate resume quality score based on content structure and completeness."""
        quality_factors = []
        
        # Length factor
        word_count = len(resume_text.split())
        if word_count >= 200:
            quality_factors.append(10)
        elif word_count >= 100:
            quality_factors.append(5)
        
        # Section completeness
        sections = ['experience', 'education', 'skills', 'summary', 'objective']
        sections_found = sum(1 for section in sections if section in resume_text.lower())
        quality_factors.append(sections_found * 4)
        
        # Contact information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
        
        if re.search(email_pattern, resume_text):
            quality_factors.append(5)
        if re.search(phone_pattern, resume_text):
            quality_factors.append(5)
        
        quality_bonus = sum(quality_factors)
        return min(100, base_score + quality_bonus)
    
    def _calculate_skills_match_score(self, skills: List[str], job_description: str, base_score: int) -> int:
        """Calculate skills match score."""
        if not skills:
            return max(30, base_score - 20)
        
        job_lower = job_description.lower()
        matched_skills = [skill for skill in skills if skill.lower() in job_lower]
        
        if len(skills) == 0:
            match_ratio = 0
        else:
            match_ratio = len(matched_skills) / len(skills)
        
        # Boost score based on skill matches
        skill_boost = int(match_ratio * 30)
        return min(100, base_score + skill_boost)
    
    def _calculate_experience_match_score(
        self, 
        experience_years: str, 
        education_level: str, 
        job_description: str, 
        base_score: int
    ) -> int:
        """Calculate experience match score."""
        score = base_score
        
        # Extract required years from job description
        job_lower = job_description.lower()
        required_years_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*years?\s*in',
            r'minimum\s*(\d+)\s*years?'
        ]
        
        required_years = 0
        for pattern in required_years_patterns:
            matches = re.findall(pattern, job_lower)
            if matches:
                required_years = max(int(match) for match in matches)
                break
        
        # Parse candidate experience
        candidate_years = 0
        if experience_years != "Not specified":
            years_match = re.search(r'(\d+)', experience_years)
            if years_match:
                candidate_years = int(years_match.group(1))
        
        # Calculate experience score
        if required_years > 0:
            if candidate_years >= required_years:
                score += 15  # Meets requirement
            elif candidate_years >= required_years * 0.7:
                score += 10  # Close to requirement
            elif candidate_years >= required_years * 0.5:
                score += 5   # Somewhat close
            else:
                score -= 10  # Significant gap
        
        # Education compensation
        if education_level in ['Masters', 'Phd']:
            score += 10
        elif education_level == 'Bachelors':
            score += 5
        
        return min(100, max(20, score))
    
    def _identify_strengths_hf(
        self, 
        skills: List[str], 
        experience_years: str, 
        education_level: str, 
        industry_keywords: List[str],
        similarity_score: float
    ) -> List[str]:
        """Identify candidate strengths based on HuggingFace analysis."""
        strengths = []
        
        if similarity_score >= 0.8:
            strengths.append("Excellent semantic match with job requirements")
        elif similarity_score >= 0.6:
            strengths.append("Strong alignment with job description")
        
        if len(skills) >= 10:
            strengths.append(f"Extensive skill set with {len(skills)} identified skills")
        elif len(skills) >= 5:
            strengths.append("Good range of technical skills")
        
        if experience_years != "Not specified":
            years_match = re.search(r'(\d+)', experience_years)
            if years_match and int(years_match.group(1)) >= 5:
                strengths.append(f"Experienced professional with {experience_years}")
        
        if education_level in ['Masters', 'Phd']:
            strengths.append(f"Advanced education: {education_level}")
        
        if len(industry_keywords) >= 3:
            strengths.append("Strong industry knowledge and terminology")
        
        return strengths[:4]  # Limit to top 4 strengths
    
    def _identify_weaknesses_hf(
        self, 
        skills: List[str], 
        experience_years: str, 
        education_level: str,
        similarity_score: float
    ) -> List[str]:
        """Identify potential candidate weaknesses."""
        weaknesses = []
        
        if similarity_score < 0.4:
            weaknesses.append("Low semantic similarity with job requirements")
        elif similarity_score < 0.6:
            weaknesses.append("Moderate alignment with job description needs improvement")
        
        if len(skills) < 3:
            weaknesses.append("Limited technical skills identified")
        
        if experience_years == "Not specified":
            weaknesses.append("Work experience duration unclear")
        
        if education_level == "Not specified":
            weaknesses.append("Educational background not clearly stated")
        
        return weaknesses[:3]  # Limit to top 3 weaknesses
    
    def _generate_ai_insights_hf(
        self, 
        overall_score: int, 
        similarity_score: float, 
        skills: List[str], 
        experience_years: str
    ) -> List[str]:
        """Generate AI-powered insights based on HuggingFace analysis."""
        insights = []
        
        # Overall assessment
        if overall_score >= 80:
            insights.append("Strong candidate with excellent fit for the role")
        elif overall_score >= 65:
            insights.append("Good candidate worth considering for interview")
        elif overall_score >= 50:
            insights.append("Moderate fit - review carefully against specific requirements")
        else:
            insights.append("Limited alignment with job requirements")
        
        # Semantic analysis insight
        if similarity_score >= 0.7:
            insights.append("Resume content shows strong semantic alignment with job posting")
        elif similarity_score >= 0.5:
            insights.append("Decent content match but could benefit from more specific examples")
        else:
            insights.append("Resume content needs better alignment with job requirements")
        
        # Skills insight
        if len(skills) >= 8:
            insights.append("Comprehensive skill set demonstrates versatility")
        elif len(skills) >= 4:
            insights.append("Good technical foundation with room for specialization")
        else:
            insights.append("Limited skill information - consider requesting additional details")
        
        return insights[:3]  # Limit to top 3 insights
    
    def _score_to_action(self, score: int) -> str:
        """Convert overall score to recommended action."""
        if score >= 80:
            return "hire"
        elif score >= 65:
            return "interview"
        elif score >= 50:
            return "maybe"
        else:
            return "reject"
    
    def _calculate_enhanced_resume_score(self, resume_text: str, job_description: str, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Enhanced fallback scoring method using comprehensive text analysis."""
        resume_lower = resume_text.lower()
        job_lower = job_description.lower()
        
        # Get extracted skills
        candidate_skills = set(skill.lower() for skill in extracted_data.get('skills', []))
        
        # Extract job required skills (more comprehensive)
        job_skills = set(self._extract_skills_from_text(job_description))
        job_skills_lower = set(skill.lower() for skill in job_skills)
        
        # Skills matching calculation (more realistic)
        if job_skills_lower:
            skills_overlap = len(candidate_skills.intersection(job_skills_lower))
            total_candidate_skills = len(candidate_skills)
            
            # More conservative job match calculation
            job_match_ratio = skills_overlap / len(job_skills_lower) if len(job_skills_lower) > 0 else 0
            skills_abundance = min(0.6, total_candidate_skills / 15)  # Reduced bonus for many skills
            
            # More realistic weighted score (75% job match, 25% skills abundance)
            base_score = (job_match_ratio * 75 + skills_abundance * 25)
            
            # Apply diminishing returns to prevent unrealistic high scores
            if base_score >= 80:
                skills_match_score = 80 + (base_score - 80) * 0.3
            elif base_score >= 60:
                skills_match_score = 60 + (base_score - 60) * 0.7
            else:
                skills_match_score = base_score * 0.9
            
            skills_match_score = min(95, skills_match_score)
        else:
            # When no job skills detected, be more conservative
            skills_match_score = min(70, len(candidate_skills) * 6 + 25)
        
        # Experience matching (more nuanced)
        experience_years = extracted_data.get('experience_years', '0 years')
        exp_number = int(re.search(r'(\d+)', experience_years).group(1)) if re.search(r'(\d+)', experience_years) else 0
        
        # Look for experience requirements in job description
        job_exp_matches = re.findall(r'(\d+)\+?\s*years?', job_lower)
        required_years = max([int(match) for match in job_exp_matches]) if job_exp_matches else 3
        
        # More differentiated and realistic experience scoring
        if exp_number == 0:
            experience_match_score = 20  # New graduate/entry level
        elif exp_number < required_years * 0.5:
            experience_match_score = 25 + (exp_number / required_years) * 20  # Significantly under
        elif exp_number < required_years * 0.8:
            experience_match_score = 45 + ((exp_number - required_years * 0.5) / (required_years * 0.3)) * 15  # Moderately under
        elif exp_number < required_years:
            experience_match_score = 60 + ((exp_number - required_years * 0.8) / (required_years * 0.2)) * 10  # Slightly under
        elif exp_number == required_years:
            experience_match_score = 75  # Exactly meets requirement
        elif exp_number <= required_years + 2:
            experience_match_score = 75 + ((exp_number - required_years) / 2) * 10  # Moderately exceeds
        else:
            # Senior level - diminishing returns after +2 years
            base_senior = 85
            extra_years = exp_number - (required_years + 2)
            senior_bonus = min(10, extra_years * 2)  # Max 10 additional points
            experience_match_score = base_senior + senior_bonus
        
        experience_match_score = min(95, max(15, experience_match_score))  # Cap between 15-95
        
        # Resume quality score
        resume_quality_score = self._calculate_resume_quality(resume_text, extracted_data)
        
        # Overall score calculation (weighted) with final calibration
        weighted_score = (
            (skills_match_score * 0.4) + 
            (experience_match_score * 0.35) + 
            (resume_quality_score * 0.25)
        )
        
        # Apply final calibration to ensure realistic score distribution
        if weighted_score > 85:
            overall_score = int(85 + (weighted_score - 85) * 0.4)  # Compress very high scores
        elif weighted_score > 70:
            overall_score = int(70 + (weighted_score - 70) * 0.7)  # Moderate compression
        else:
            overall_score = int(weighted_score * 0.95)  # Slight compression for lower scores
            
        overall_score = min(95, max(20, overall_score))  # Final bounds
        
        # Generate insights
        strengths, weaknesses, insights = self._generate_detailed_insights(
            candidate_skills, job_skills_lower, exp_number, required_years, 
            extracted_data, overall_score
        )
        
        return {
            "resumeScore": int(resume_quality_score),
            "skillsMatch": int(skills_match_score),
            "experienceMatch": int(experience_match_score),
            "overallScore": overall_score,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "aiInsights": insights,
            "confidenceLevel": 75,
            "recommendedAction": self._get_recommendation(overall_score)
        }
    
    def _calculate_resume_quality(self, resume_text: str, extracted_data: Dict[str, Any]) -> float:
        """Calculate resume quality based on completeness and structure."""
        score = 50  # More conservative base score
        
        # Length and completeness (more balanced scoring)
        text_length = len(resume_text)
        if text_length > 2000:
            score += 20  # Comprehensive resume
        elif text_length > 1000:
            score += 15  # Good detail level
        elif text_length > 500:
            score += 10  # Adequate detail
        elif text_length > 200:
            score += 5   # Minimal detail
        # Below 200 chars gets no bonus
        
        # Skills presence and quantity
        skills = extracted_data.get('skills', [])
        if len(skills) >= 10:
            score += 15  # Many relevant skills
        elif len(skills) >= 5:
            score += 12  # Good skill set
        elif len(skills) >= 3:
            score += 8   # Some skills
        elif len(skills) >= 1:
            score += 5   # Few skills
        
        # Education information (more conservative)
        education = extracted_data.get('education_level', '').lower()
        if 'phd' in education or 'doctorate' in education:
            score += 12
        elif 'master' in education or 'mba' in education:
            score += 10
        elif 'bachelor' in education:
            score += 8
        elif 'associate' in education:
            score += 4
        # No education specified gets no penalty
        
        # Professional sections
        resume_lower = resume_text.lower()
        sections = ['experience', 'skills', 'education', 'projects', 'certifications']
        found_sections = sum(1 for section in sections if section in resume_lower)
        score += found_sections * 2  # Reduced bonus per section
        
        # Keywords indicating achievements (more selective)
        achievement_keywords = ['achieved', 'improved', 'increased', 'reduced', 'led', 'managed', 'developed', 'implemented', 'created', 'built']
        achievements = sum(1 for keyword in achievement_keywords if keyword in resume_lower)
        score += min(achievements, 8)  # Max 8 points for achievements
        
        # Professional formatting indicators
        professional_indicators = ['email', 'phone', 'linkedin', 'github']
        found_indicators = sum(1 for indicator in professional_indicators if indicator in resume_lower)
        score += found_indicators * 2
        
        return min(95, max(35, score))  # Score between 35-95
    
    def _generate_detailed_insights(self, candidate_skills: set, job_skills: set, 
                                  exp_years: int, required_years: int, 
                                  extracted_data: Dict[str, Any], overall_score: int) -> tuple:
        """Generate detailed strengths, weaknesses, and insights."""
        strengths = []
        weaknesses = []
        insights = []
        
        # Skills analysis
        matching_skills = candidate_skills.intersection(job_skills)
        missing_skills = job_skills - candidate_skills
        
        if len(matching_skills) > 3:
            strengths.append(f"Strong technical skill alignment ({len(matching_skills)} matching skills)")
        elif len(matching_skills) > 1:
            strengths.append("Good technical foundation with relevant skills")
        
        if missing_skills and len(missing_skills) <= 3:
            weaknesses.append(f"Missing key skills: {', '.join(list(missing_skills)[:3])}")
        elif len(missing_skills) > 3:
            weaknesses.append("Several important technical skills not demonstrated")
        
        # Experience analysis
        if exp_years >= required_years:
            strengths.append(f"Meets experience requirement ({exp_years}+ years)")
        elif exp_years >= required_years * 0.8:
            strengths.append("Nearly meets experience requirement")
        else:
            weaknesses.append(f"Below preferred experience level (has {exp_years}, needs {required_years}+)")
        
        # Education
        education = extracted_data.get('education_level', '').lower()
        if 'phd' in education or 'master' in education:
            strengths.append("Advanced education background")
        elif 'bachelor' in education:
            strengths.append("Solid educational foundation")
        
        # Generate insights based on overall score
        if overall_score >= 85:
            insights.append("Excellent candidate - strong match for the role")
            insights.append("Recommend priority consideration for interview")
        elif overall_score >= 70:
            insights.append("Good candidate with potential for the role")
            insights.append("Consider for interview with focus on skill gaps")
        elif overall_score >= 55:
            insights.append("Marginal fit - may need additional training or experience")
            insights.append("Evaluate cultural fit and growth potential")
        else:
            insights.append("Below minimum requirements for this role")
            insights.append("Consider for more junior positions if available")
        
        return strengths[:3], weaknesses[:2], insights[:3]
    
    def _get_recommendation(self, score: int) -> str:
        """Get hiring recommendation based on calibrated score."""
        if score >= 80:
            return "hire"
        elif score >= 65:
            return "interview"
        elif score >= 50:
            return "maybe"
        else:
            return "reject"
    
    def _calculate_basic_resume_score(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """Fallback scoring method using basic text analysis."""
        # Basic keyword matching and scoring
        resume_lower = resume_text.lower()
        job_lower = job_description.lower()
        
        # Load keywords dynamically from configuration
        skills_db = config_manager.get_skill_variations()
        name_patterns = config_manager.get_name_patterns()
        
        # Flatten tech keywords from skills database
        tech_keywords = []
        for skill_category in skills_db.values():
            if isinstance(skill_category, dict):
                for variations in skill_category.values():
                    if isinstance(variations, list):
                        tech_keywords.extend(variations)
        
        experience_keywords = ['years', 'experience', 'developed', 'managed', 'led', 'implemented', 'designed', 'worked', 'created', 'built', 'engineered']
        education_keywords = name_patterns.get('education_keywords', ['degree', 'bachelor', 'master', 'university', 'college', 'certification', 'diploma', 'phd', 'certifications'])
        
        # Calculate basic matches
        tech_matches = sum(1 for keyword in tech_keywords if keyword in resume_lower and keyword in job_lower)
        exp_matches = sum(1 for keyword in experience_keywords if keyword in resume_lower)
        edu_matches = sum(1 for keyword in education_keywords if keyword in resume_lower)
        
        # Calculate scores (0-100)
        tech_score = min(100, (tech_matches / max(1, len(tech_keywords))) * 100 + 50)
        exp_score = min(100, (exp_matches / max(1, len(experience_keywords))) * 100 + 60)
        resume_score = min(100, 60 + (len(resume_text) / 50))  # Length-based scoring
        overall_score = (tech_score + exp_score + resume_score) / 3
        
        return {
            "resumeScore": int(resume_score),
            "skillsMatch": int(tech_score),
            "experienceMatch": int(exp_score),
            "overallScore": int(overall_score),
            "strengths": ["Relevant background identified", "Technical skills present", "Professional experience evident"],
            "weaknesses": ["Could benefit from more specific examples", "Additional certifications recommended"],
            "aiInsights": ["Basic analysis completed", "Consider detailed technical interview", "Profile shows potential for role"]
        }

    async def generate_team_insights(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Generate AI-powered insights for team members."""
        try:
            employee = context.get("employee", {})
            performance_score = employee.get("performanceScore", 75)
            role = employee.get("role", "")
            department = employee.get("department", "")
            skills = employee.get("skills", [])
            
            # Generate insights using AI model if available
            if self.gemini_model:
                prompt = f"""
                Analyze this team member profile and provide insights:
                
                Name: {employee.get('name', 'Team Member')}
                Role: {role}
                Department: {department}
                Skills: {', '.join(skills[:5]) if skills else 'Not specified'}
                Performance Score: {performance_score}%
                Career Goals: {employee.get('careerGoals', 'Not specified')}
                
                Provide structured insights in the following format:
                - Strengths: 2-3 key strengths
                - Growth Areas: 2-3 areas for development
                - Recommendations: 2-3 actionable recommendations
                - Risk Level: low, medium, or high
                
                Keep it professional and constructive.
                """
                
                try:
                    response = await asyncio.to_thread(
                        self.gemini_model.generate_content,
                        prompt
                    )
                    
                    # Parse AI response
                    insights_text = response.text.strip()
                    return self._parse_team_insights(insights_text, performance_score)
                    
                except Exception as e:
                    logger.warning(f"AI insights generation failed, using fallback: {e}")
                    return self._generate_fallback_team_insights(employee, performance_score)
            
            else:
                return self._generate_fallback_team_insights(employee, performance_score)
                
        except Exception as e:
            logger.error(f"Error generating team insights: {e}")
            return self._generate_fallback_team_insights(context.get("employee", {}), 75)
    
    def _parse_team_insights(self, insights_text: str, performance_score: float) -> Dict[str, Any]:
        """Parse AI-generated insights text into structured format."""
        try:
            # Extract sections using regex
            strengths_match = re.search(r'strengths?:?\s*(.+?)(?=growth|recommendations|risk|$)', insights_text, re.IGNORECASE | re.DOTALL)
            growth_match = re.search(r'growth areas?:?\s*(.+?)(?=recommendations|risk|$)', insights_text, re.IGNORECASE | re.DOTALL)
            recommendations_match = re.search(r'recommendations?:?\s*(.+?)(?=risk|$)', insights_text, re.IGNORECASE | re.DOTALL)
            risk_match = re.search(r'risk level:?\s*(\w+)', insights_text, re.IGNORECASE)
            
            # Parse bullet points
            def parse_bullets(text):
                if not text:
                    return []
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                bullets = []
                for line in lines:
                    # Remove bullet markers
                    clean_line = re.sub(r'^[-*•]\s*', '', line.strip())
                    if clean_line and len(clean_line) > 10:
                        bullets.append(clean_line)
                return bullets[:3]  # Limit to 3 items
            
            strengths = parse_bullets(strengths_match.group(1) if strengths_match else "")
            growth_areas = parse_bullets(growth_match.group(1) if growth_match else "")
            recommendations = parse_bullets(recommendations_match.group(1) if recommendations_match else "")
            
            risk_level = "low"
            if risk_match:
                risk_text = risk_match.group(1).lower()
                if "high" in risk_text:
                    risk_level = "high"
                elif "medium" in risk_text:
                    risk_level = "medium"
            
            return {
                "strengths": strengths or ["Strong technical capabilities", "Reliable team member"],
                "growthAreas": growth_areas or ["Leadership development", "Cross-functional collaboration"],
                "recommendations": recommendations or ["Consider mentoring opportunities", "Explore advanced training"],
                "riskLevel": risk_level
            }
            
        except Exception as e:
            logger.warning(f"Error parsing AI insights, using fallback: {e}")
            return self._generate_fallback_team_insights({"performanceScore": performance_score}, performance_score)
    
    def _generate_fallback_team_insights(self, employee: Dict[str, Any], performance_score: float) -> Dict[str, Any]:
        """Generate fallback insights when AI is not available."""
        role = employee.get("role", "").lower()
        department = employee.get("department", "").lower()
        
        # Role-based strengths
        strengths = []
        if "senior" in role:
            strengths.append("Experienced professional with senior-level expertise")
        if "manager" in role or "lead" in role:
            strengths.append("Leadership and team management capabilities")
        if "engineer" in role:
            strengths.append("Strong technical and problem-solving skills")
        if "designer" in role:
            strengths.append("Creative and user-focused design thinking")
        
        # Performance-based insights
        if performance_score >= 90:
            strengths.append("Exceptional performance and high impact contributor")
        elif performance_score >= 80:
            strengths.append("Consistently high performer and reliable team member")
        elif performance_score >= 70:
            strengths.append("Solid contributor with good work quality")
        
        # Default strengths if none identified
        if not strengths:
            strengths = ["Dedicated team member", "Professional approach to work"]
        
        # Department-based growth areas
        growth_areas = []
        if department in ["engineering", "tech", "development"]:
            growth_areas.extend(["Cross-functional collaboration", "Technical leadership"])
        elif department in ["design", "creative"]:
            growth_areas.extend(["Strategic design thinking", "User research skills"])
        elif department in ["marketing", "sales"]:
            growth_areas.extend(["Data-driven decision making", "Digital marketing trends"])
        else:
            growth_areas.extend(["Leadership development", "Process optimization"])
        
        # Performance-based recommendations
        recommendations = []
        if performance_score >= 85:
            recommendations.extend(["Consider for leadership opportunities", "Mentoring junior team members"])
        elif performance_score >= 70:
            recommendations.extend(["Focus on skill specialization", "Increase project ownership"])
        else:
            recommendations.extend(["Performance improvement plan", "Additional training and support"])
        
        # Default recommendations
        if not recommendations:
            recommendations = ["Regular 1:1 check-ins", "Professional development opportunities"]
        
        # Risk assessment
        risk_level = "low"
        if performance_score < 60:
            risk_level = "high"
        elif performance_score < 75:
            risk_level = "medium"
        
        return {
            "strengths": strengths[:3],
            "growthAreas": growth_areas[:3],
            "recommendations": recommendations[:3],
            "riskLevel": risk_level
        }

# Global AI service instance
ai_service = AIService()